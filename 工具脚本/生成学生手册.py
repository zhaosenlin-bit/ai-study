from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "01-classroom-student-handbook.md"
OUTPUT = ROOT / "docs" / "presentation" / "ai-study-课堂展示讲义-v0.3.docx"


FONT = "Microsoft YaHei"
HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
TABLE_HEADER = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, size: int | float | None = None, bold: bool = False, color: RGBColor | None = None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_bg(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, HEADING_BLUE, 18, 10),
        ("Heading 2", 13, HEADING_BLUE, 14, 7),
        ("Heading 3", 12, HEADING_DARK, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.2


def add_cover(doc: Document):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(36)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("ai-study")
    set_run_font(run, 30, True, HEADING_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("三科自适应伴学 Agent 课堂展示讲义")
    set_run_font(run, 16, True, HEADING_DARK)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("2026 iFLYTEK AI 开发者大赛 · Adaptive-LPDS · 2026-08-18")
    set_run_font(run, 10.5, False, MUTED)

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [2200, 7160])
    rows = [
        ("项目定位", "小学 3-6 年级语文、数学、英语三科 AI 学习伙伴"),
        ("核心闭环", "诊断 -> 路径规划 -> 脚手架辅导 -> 错题复习 -> 家长报告"),
        ("统一技术", "React + Vite + TypeScript，FastAPI，Agent，知识图谱，可插拔模型"),
        ("团队分工", "A 后端 Agent，B 数学，C 语文，D 英语，E 前端联调演示"),
        ("课堂目标", "让每位同学知道自己负责什么、先做什么、怎么提交、怎么验收"),
    ]
    for idx, (label, value) in enumerate(rows):
        cells = table.rows[idx].cells
        cells[0].text = label
        cells[1].text = value
        set_cell_bg(cells[0], TABLE_HEADER)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, 10, bold=(cell is cells[0]))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("先跑通完整闭环，再打磨亮点；先统一接口，再并行开发。")
    set_run_font(run, 12, True, HEADING_DARK)

    doc.add_page_break()


def add_footer(doc: Document):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("ai-study 项目课堂展示讲义")
        set_run_font(run, 8.5, False, MUTED)


def add_paragraph_with_inline_code(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6 if style is None else 4)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, 10, False, RGBColor(0x9B, 0x1C, 0x1C))
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        else:
            run = p.add_run(part)
            set_run_font(run, 10.5)
    return p


def add_code_block(doc: Document, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(9)


def parse_table(block: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in block:
        stripped = line.strip().strip("|")
        cells = [cell.strip() for cell in stripped.split("|")]
        rows.append(cells)
    header = rows[0]
    body = [row for row in rows[2:] if not all(set(cell.replace(":", "").replace("-", "").strip()) <= set() for cell in row)]
    return header, body


def add_markdown_table(doc: Document, block: list[str]):
    header, body = parse_table(block)
    cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [9360 // cols] * cols
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)

    for c, label in enumerate(header):
        cell = table.rows[0].cells[c]
        set_cell_bg(cell, TABLE_HEADER)
        cell.text = label
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, 9.5, True, HEADING_DARK)

    for r, row in enumerate(body, start=1):
        for c in range(cols):
            value = row[c] if c < len(row) else ""
            cell = table.rows[r].cells[c]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, 9)

    doc.add_paragraph()


def add_check_item(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.replace("[ ]", "☐").strip())
    set_run_font(run, 10)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    source = SOURCE.read_text(encoding="utf-8").splitlines()

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

    for raw in source[1:]:
        line = raw.rstrip()

        if line.startswith("```"):
            flush_table()
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|"):
            table_lines.append(line)
            continue
        flush_table()

        if not line.strip():
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- [ ]"):
            add_check_item(doc, line[2:])
        elif line.startswith("- "):
            add_paragraph_with_inline_code(doc, line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            add_paragraph_with_inline_code(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            add_paragraph_with_inline_code(doc, line)

    flush_table()
    add_footer(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()